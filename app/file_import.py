from __future__ import annotations

import csv
import html
import io
import multiprocessing
import os
import posixpath
import re
import sqlite3
import tempfile
import time
import zipfile
from html.parser import HTMLParser
from pathlib import PurePosixPath
from xml.etree import ElementTree


class ImportFileError(ValueError):
    """文件导入业务错误；不要命名为 ImportError，避免遮蔽内置异常。"""


_ARCHIVE_UNCOMPRESSED_LIMIT = 64 * 1024 * 1024
_MAX_ARCHIVE_MEMBERS = 5000
_MAX_PDF_PAGES = 5000


class _ArchiveReader:
    """按“实际解压字节数”封顶读取 zip 成员。

    zip 头部声明的 file_size 可被伪造，因此不能用声明值做限额：
    这里流式读取并累计实际解压出的字节数，超限立即中止，
    防止伪造元数据的解压炸弹把整块内容读入内存。
    """

    def __init__(self, archive: zipfile.ZipFile, limit: int):
        self._archive = archive
        self._limit = limit
        self.total = 0

    def read(self, member) -> bytes:
        buffer = io.BytesIO()
        with self._archive.open(member) as source:
            while True:
                chunk = source.read(64 * 1024)
                if not chunk:
                    break
                self.total += len(chunk)
                if self.total > self._limit:
                    raise ImportFileError(
                        f"解压后过大，最多支持 {self._limit} 字节"
                    )
                buffer.write(chunk)
        return buffer.getvalue()


def _reject_xml_entity_declarations(data: bytes, label: str) -> None:
    """拒绝 DOCTYPE/ENTITY 声明，防止标准库 XML 解析器做实体扩展放大。"""
    if b"<!DOCTYPE" in data or b"<!ENTITY" in data:
        raise ImportFileError(f"{label} 包含不安全的 XML 实体声明")


def _check_budget(size: int, max_chars: int) -> None:
    if size > max_chars:
        raise ImportFileError(f"正文过长，最多支持 {max_chars} 个字符")


def _decode_text(data: bytes) -> str:
    candidates = ["utf-8-sig", "gb18030"]
    try:
        import chardet

        detected = str((chardet.detect(data) or {}).get("encoding") or "").strip()
        if detected:
            candidates.append(detected)
    except ImportError:
        pass
    candidates.append("latin-1")
    for encoding in dict.fromkeys(candidates):
        try:
            return data.decode(encoding)
        except (LookupError, UnicodeDecodeError):
            continue
    raise ImportFileError("无法识别文本文件编码")


def _validate_office_archive(data: bytes, file_type: str) -> None:
    try:
        with zipfile.ZipFile(io.BytesIO(data)) as archive:
            infos = archive.infolist()
            if len(infos) > _MAX_ARCHIVE_MEMBERS:
                raise ImportFileError(f"{file_type} 内文件数量过多")
            if sum(info.file_size for info in infos) > _ARCHIVE_UNCOMPRESSED_LIMIT:
                raise ImportFileError(f"{file_type} 解压后过大")
            reader = _ArchiveReader(archive, _ARCHIVE_UNCOMPRESSED_LIMIT)
            for info in infos:
                _safe_zip_path(info.filename)
                if info.is_dir():
                    continue
                name = info.filename.lower()
                if name.endswith((".xml", ".rels")) or "/_rels/" in info.filename:
                    _reject_xml_entity_declarations(
                        reader.read(info), f"{file_type} XML"
                    )
    except (zipfile.BadZipFile, RuntimeError, OSError) as exc:
        raise ImportFileError(f"{file_type} 文件损坏或格式不正确") from exc


class _TextExtractor(HTMLParser):
    _BLOCKS = {
        "address", "article", "aside", "blockquote", "br", "div", "figcaption",
        "figure", "footer", "h1", "h2", "h3", "h4", "h5", "h6", "header",
        "hr", "li", "main", "nav", "ol", "p", "pre", "section", "table",
        "tr", "ul",
    }
    _SKIP = {"script", "style", "svg"}

    def __init__(self):
        super().__init__(convert_charrefs=True)
        self.parts: list[str] = []
        self.skip_depth = 0

    def handle_starttag(self, tag: str, attrs) -> None:
        tag = tag.lower()
        if tag in self._SKIP:
            self.skip_depth += 1
        elif tag in self._BLOCKS and not self.skip_depth:
            self.parts.append("\n")

    def handle_endtag(self, tag: str) -> None:
        tag = tag.lower()
        if tag in self._SKIP and self.skip_depth:
            self.skip_depth -= 1
        elif tag in self._BLOCKS and not self.skip_depth:
            self.parts.append("\n")

    def handle_data(self, data: str) -> None:
        if not self.skip_depth:
            self.parts.append(data)

    def text(self) -> str:
        raw = "".join(self.parts).replace("\xa0", " ")
        raw = re.sub(r"[ \t\r\f\v]+", " ", raw)
        raw = re.sub(r" *\n *", "\n", raw)
        raw = re.sub(r"\n{3,}", "\n\n", raw)
        return raw.strip()


_PARAGRAPH_BLOCK_KINDS = {
    "h1": "h1", "h2": "h2", "h3": "h3", "h4": "h4", "h5": "h5", "h6": "h6",
    "p": "p", "blockquote": "quote", "li": "li",
}
_PARAGRAPH_INLINE_ALLOW = {"em", "i", "strong", "b", "br"}


class _ParagraphBuilder(HTMLParser):
    """把 XHTML 正文拆成安全段落：只保留基础行内格式，丢弃属性和危险标签。"""

    def __init__(self):
        super().__init__(convert_charrefs=True)
        self.paragraphs: list[dict] = []
        self._kind = "p"
        self._buf: list[str] = []
        self._stack: list[str] = []
        self._skip_depth = 0

    def handle_starttag(self, tag: str, attrs) -> None:
        tag = tag.lower()
        if tag in _TextExtractor._SKIP:
            self._skip_depth += 1
            return
        if self._skip_depth:
            return
        if tag in _PARAGRAPH_BLOCK_KINDS:
            self._flush()
            self._stack.append(tag)
            self._kind = _PARAGRAPH_BLOCK_KINDS[tag]
        elif tag in _PARAGRAPH_INLINE_ALLOW:
            if tag == "br":
                self._buf.append("<br>")
            else:
                self._buf.append(f"<{tag}>")

    def handle_endtag(self, tag: str) -> None:
        tag = tag.lower()
        if tag in _TextExtractor._SKIP and self._skip_depth:
            self._skip_depth -= 1
            return
        if self._skip_depth:
            return
        if tag in _PARAGRAPH_INLINE_ALLOW and tag != "br":
            self._buf.append(f"</{tag}>")
        elif tag in _PARAGRAPH_BLOCK_KINDS:
            self._flush()
            if self._stack:
                self._stack.pop()
            self._kind = self._stack[-1] if self._stack else "p"

    def handle_data(self, data: str) -> None:
        if not self._skip_depth:
            self._buf.append(html.escape(data))

    def _flush(self) -> None:
        fragment = "".join(self._buf).strip()
        self._buf = []
        if not fragment:
            return
        plain = html.unescape(re.sub(r"<[^>]+>", "", fragment))
        plain = re.sub(r"[ \t\r\f\v]+", " ", plain).strip()
        self.paragraphs.append(
            {"kind": self._kind, "html": fragment, "text": plain}
        )

    def result(self) -> list[dict]:
        self._flush()
        return self.paragraphs


def _html_paragraphs(raw: bytes) -> list[dict]:
    parser = _ParagraphBuilder()
    parser.feed(raw.decode("utf-8", errors="replace"))
    return parser.result()


def _paragraphs_from_text(text: str) -> list[dict]:
    blocks = [block.strip() for block in re.split(r"\n\s*\n", text) if block.strip()]
    paragraphs = []
    for block in blocks:
        plain = re.sub(r"[ \t\r\f\v]+", " ", block).strip()
        paragraphs.append(
            {
                "kind": "p",
                "html": "<br>".join(html.escape(line) for line in block.splitlines()),
                "text": plain,
            }
        )
    return paragraphs


def _safe_zip_path(path: str) -> str:
    normalized = posixpath.normpath(path.replace("\\", "/"))
    pure = PurePosixPath(normalized)
    if normalized.startswith("/") or ".." in pure.parts:
        raise ImportFileError("EPUB 内包含不安全的文件路径")
    return normalized


def _html_to_text(raw: bytes) -> str:
    parser = _TextExtractor()
    parser.feed(raw.decode("utf-8", errors="replace"))
    return parser.text()


def _chapter_title(raw: bytes, fallback: str) -> str:
    html = raw.decode("utf-8", errors="replace")
    match = re.search(
        r"<(?:h1|h2|title)\b[^>]*>(.*?)</(?:h1|h2|title)>",
        html,
        flags=re.IGNORECASE | re.DOTALL,
    )
    if match:
        title = _html_to_text(match.group(1).encode()).splitlines()[0].strip()
        if title:
            return title[:300]
    return fallback[:300]


def _epub_navigation_titles(
    reader: _ArchiveReader,
    names: dict[str, str],
    opf,
    opf_path: str,
) -> dict[str, str]:
    """读取 EPUB2/3 导航目录中的具体章节名称，并映射到正文路径。"""
    base = posixpath.dirname(opf_path)
    titles: dict[str, str] = {}
    for item in opf.findall(".//{*}manifest/{*}item"):
        href = item.attrib.get("href", "")
        if not href:
            continue
        properties = item.attrib.get("properties", "").split()
        media_type = item.attrib.get("media-type", "")
        path = _safe_zip_path(posixpath.join(base, href.split("#", 1)[0]))
        if path not in names:
            continue
        if "nav" in properties:
            raw = reader.read(names[path]).decode("utf-8", errors="replace")
            for match in re.finditer(
                r"<a\b[^>]*\bhref\s*=\s*['\"]([^'\"]+)['\"][^>]*>(.*?)</a>",
                raw,
                flags=re.IGNORECASE | re.DOTALL,
            ):
                target = _safe_zip_path(
                    posixpath.join(posixpath.dirname(path), match.group(1).split("#", 1)[0])
                )
                label = _html_to_text(match.group(2).encode()).strip()
                if target in names and label:
                    titles.setdefault(target, label[:300])
        elif media_type == "application/x-dtbncx+xml":
            ncx_raw = reader.read(names[path])
            _reject_xml_entity_declarations(ncx_raw, "EPUB 导航")
            try:
                ncx = ElementTree.fromstring(ncx_raw)
            except ElementTree.ParseError:
                continue
            for point in ncx.findall(".//{*}navPoint"):
                content = point.find("./{*}content")
                label = point.find("./{*}navLabel/{*}text")
                if content is None or label is None or not (label.text or "").strip():
                    continue
                target = _safe_zip_path(
                    posixpath.join(
                        posixpath.dirname(path),
                        content.attrib.get("src", "").split("#", 1)[0],
                    )
                )
                if target in names:
                    titles.setdefault(target, (label.text or "").strip()[:300])
    return titles


def extract_epub_chapters(
    data: bytes,
    max_chapter_bytes: int = _ARCHIVE_UNCOMPRESSED_LIMIT,
    max_chars: int | None = None,
) -> list[dict]:
    try:
        archive = zipfile.ZipFile(io.BytesIO(data))
    except zipfile.BadZipFile as exc:
        raise ImportFileError("EPUB 文件损坏或格式不正确") from exc
    with archive:
        try:
            infos = archive.infolist()
            if len(infos) > 2000:
                raise ImportFileError("EPUB 内文件数量过多")
            if sum(info.file_size for info in infos) > _ARCHIVE_UNCOMPRESSED_LIMIT:
                raise ImportFileError("EPUB 解压后过大")
            names = {_safe_zip_path(info.filename): info.filename for info in infos}
            metadata_reader = _ArchiveReader(archive, _ARCHIVE_UNCOMPRESSED_LIMIT)
            container_raw = metadata_reader.read("META-INF/container.xml")
            _reject_xml_entity_declarations(container_raw, "EPUB 目录")
            container = ElementTree.fromstring(container_raw)
            rootfile = container.find(".//{*}rootfile")
            opf_path = _safe_zip_path(rootfile.attrib["full-path"] if rootfile is not None else "")
            if not opf_path:
                raise KeyError
            opf_raw = metadata_reader.read(names[opf_path])
            _reject_xml_entity_declarations(opf_raw, "EPUB 目录")
            opf = ElementTree.fromstring(opf_raw)
        except (KeyError, ElementTree.ParseError) as exc:
            raise ImportFileError("EPUB 缺少有效的书籍目录") from exc
        except zipfile.BadZipFile as exc:
            raise ImportFileError("EPUB 文件损坏或格式不正确") from exc

        manifest = {
            item.attrib.get("id", ""): item.attrib.get("href", "")
            for item in opf.findall(".//{*}manifest/{*}item")
        }
        navigation_titles = _epub_navigation_titles(
            metadata_reader, names, opf, opf_path
        )
        book_title_node = opf.find(".//{*}metadata/{*}title")
        book_title = (book_title_node.text or "").strip() if book_title_node is not None else ""
        base = posixpath.dirname(opf_path)
        chapter_paths = []
        for itemref in opf.findall(".//{*}spine/{*}itemref"):
            href = manifest.get(itemref.attrib.get("idref", ""), "")
            if not href:
                continue
            path = _safe_zip_path(posixpath.join(base, href.split("#", 1)[0]))
            if path in names:
                chapter_paths.append(path)
        if not chapter_paths:
            raise ImportFileError("EPUB 没有可读取的正文章节")
        chapter_infos = {
            _safe_zip_path(info.filename): info for info in infos
        }
        if sum(chapter_infos[path].file_size for path in chapter_paths) > max_chapter_bytes:
            raise ImportFileError("EPUB 正文章节解压后过大")
        chapters = []
        total_chars = 0
        chapter_reader = _ArchiveReader(archive, max_chapter_bytes)
        for index, path in enumerate(chapter_paths, start=1):
            try:
                raw = chapter_reader.read(names[path])
            except zipfile.BadZipFile as exc:
                raise ImportFileError("EPUB 文件损坏或格式不正确") from exc
            paragraphs = _html_paragraphs(raw)
            text = "\n\n".join(paragraph["text"] for paragraph in paragraphs)
            if not text:
                continue
            total_chars += len(text) + (2 if chapters else 0)
            if max_chars is not None:
                _check_budget(total_chars, max_chars)
            fallback = PurePosixPath(path).stem.replace("_", " ").replace("-", " ")
            title = navigation_titles.get(path) or _chapter_title(
                raw, fallback or f"第 {index} 章"
            )
            if book_title and title.casefold() == book_title.casefold():
                title = fallback or f"第 {index} 章"
            chapters.append(
                {
                    "title": title[:300],
                    "text": text,
                    "paragraphs": paragraphs,
                }
            )
        if not chapters:
            raise ImportFileError("EPUB 正文为空")
        return chapters


def extract_epub(data: bytes, max_uncompressed_bytes: int = _ARCHIVE_UNCOMPRESSED_LIMIT) -> str:
    chapters = extract_epub_chapters(
        data, max_chapter_bytes=max_uncompressed_bytes
    )
    return "\n\n".join(chapter["text"] for chapter in chapters).strip()


_CHAPTER_ROMAN = r"[ivxlcdm]{2,}"
_CHAPTER_ROMAN_ANY = r"[ivxlcdm]+"
_CH_NUM = r"[一二三四五六七八九十百千万零\d]+"
_CHAPTER_WORD_NUM = (
    r"(?:one|two|three|four|five|six|seven|eight|nine|ten|"
    r"eleven|twelve|thirteen|fourteen|fifteen|sixteen|seventeen|eighteen|"
    r"nineteen|twenty|thirty|forty|fifty|sixty|seventy|eighty|ninety|"
    r"hundred|first|second|third|fourth|fifth|sixth|seventh|eighth|ninth|tenth)"
)
_CHAPTER_NUM = rf"(?:\d+|{_CHAPTER_ROMAN_ANY}|{_CHAPTER_WORD_NUM}(?:\s+{_CHAPTER_WORD_NUM})?)"
_CHAPTER_KEYWORDS = (
    r"(?:chapter|ch|chap|part|book|section|act|unit|lesson|volume|vol|"
    r"episode|scene|stage|level|day|week|month|prologue|epilogue|"
    r"introduction|preface|foreword|afterword|preamble|interlude|postscript|"
    r"dedication|acknowledgements?|authors?\s*note)"
)
_CHAPTER_SEP = r"\s*[.:：\-–—|]*\s*"
_CHAPTER_HEADING_RE = re.compile(
    rf"^(?:"
    rf"#{{1,4}}\s+.+|"
    rf"{_CHAPTER_KEYWORDS}{_CHAPTER_SEP}(?:the\s+)?{_CHAPTER_NUM}"
    rf"(?:{_CHAPTER_SEP}.*)?|"
    rf"(?:the\s+)?{_CHAPTER_WORD_NUM}\s+(?:chapter|part|book|section|act|lesson)\s*.*|"
    rf"(?:prologue|epilogue|introduction|preface|foreword|afterword|preamble|"
    rf"interlude|postscript|dedication|acknowledgements?|authors?\s*note)"
    rf"{_CHAPTER_SEP}$|"
    rf"第{_CHAPTER_SEP}{_CH_NUM}{_CHAPTER_SEP}[章节卷回部集幕].*|"
    rf"(?:序章|序言|引子|楔子|番外|后记|卷首语|尾声|开篇|前传|后传|序幕|终章|正文).*|"
    rf"\d{{1,4}}{_CHAPTER_SEP}$|"
    rf"\d+[.)]{_CHAPTER_SEP}[^\s\d].*|"
    rf"{_CHAPTER_ROMAN}[.)]{_CHAPTER_SEP}[^\s\d].*"
    rf")$",
    flags=re.IGNORECASE,
)


def split_text_chapters(text: str) -> list[dict]:
    """按常见中英文书籍标题拆章；没有标题时仍作为单章阅读。"""
    lines = text.splitlines()
    chapters: list[dict] = []
    title = "正文"
    body: list[str] = []
    for line in lines:
        stripped = line.strip()
        if stripped and _CHAPTER_HEADING_RE.match(stripped):
            current = "\n".join(body).strip()
            if current:
                chapters.append(
                    {
                        "title": title,
                        "text": current,
                        "paragraphs": _paragraphs_from_text(current),
                    }
                )
            title = re.sub(r"^#{1,3}\s+", "", stripped)[:300]
            body = []
        else:
            body.append(line)
    current = "\n".join(body).strip()
    if current:
        chapters.append(
            {
                "title": title,
                "text": current,
                "paragraphs": _paragraphs_from_text(current),
            }
        )
    return chapters or [
        {"title": "正文", "text": text, "paragraphs": _paragraphs_from_text(text)}
    ]


def _extract_pdf(data: bytes, max_chars: int) -> tuple[str, list[dict]]:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise ImportFileError("服务器尚未安装 PDF 解析组件") from exc
    try:
        reader = PdfReader(io.BytesIO(data), strict=False)
        if len(reader.pages) > _MAX_PDF_PAGES:
            raise ImportFileError(f"PDF 页数过多，最多支持 {_MAX_PDF_PAGES} 页")
        pages = []
        total_chars = 0
        for page in reader.pages:
            value = str(page.extract_text() or "").strip()
            if not value:
                continue
            total_chars += len(value) + (2 if pages else 0)
            _check_budget(total_chars, max_chars)
            pages.append(value)
    except ImportFileError:
        raise
    except Exception as exc:
        raise ImportFileError("PDF 文件损坏、加密或无法提取文字") from exc
    text = "\n\n".join(page for page in pages if page).strip()
    return text, split_text_chapters(text)


def _extract_docx(data: bytes, max_chars: int) -> tuple[str, list[dict]]:
    _validate_office_archive(data, "DOCX")
    try:
        from docx import Document
    except ImportError as exc:
        raise ImportFileError("服务器尚未安装 DOCX 解析组件") from exc
    try:
        document = Document(io.BytesIO(data))
    except Exception as exc:
        raise ImportFileError("DOCX 文件损坏或无法读取") from exc
    chapters: list[dict] = []
    title = "正文"
    body: list[str] = []
    all_parts: list[str] = []
    total_chars = 0
    for paragraph in document.paragraphs:
        value = paragraph.text.strip()
        if not value:
            continue
        total_chars += len(value) + (2 if all_parts else 0)
        _check_budget(total_chars, max_chars)
        all_parts.append(value)
        style_name = str(getattr(paragraph.style, "name", "") or "").lower()
        if style_name.startswith("heading"):
            if body:
                chapters.append({"title": title[:300], "text": "\n".join(body)})
            title = value
            body = []
        else:
            body.append(value)
    if body:
        chapters.append({"title": title[:300], "text": "\n".join(body)})
    text = "\n\n".join(all_parts)
    return text, chapters or split_text_chapters(text)


def _extract_csv(data: bytes, max_chars: int) -> tuple[str, list[dict]]:
    decoded = _decode_text(data)
    try:
        rows = []
        total_chars = 0
        for row in csv.reader(io.StringIO(decoded)):
            value = " ".join(cell.strip() for cell in row if cell.strip())
            if value:
                total_chars += len(value) + (1 if rows else 0)
                _check_budget(total_chars, max_chars)
                rows.append(value)
    except csv.Error as exc:
        raise ImportFileError("CSV 文件格式不正确") from exc
    text = "\n".join(row for row in rows if row)
    return text, [{"title": "CSV 数据", "text": text}]


def _extract_xlsx(data: bytes, max_chars: int) -> tuple[str, list[dict]]:
    _validate_office_archive(data, "Excel")
    total_chars = 0
    try:
        from openpyxl import load_workbook
    except ImportError as exc:
        raise ImportFileError("服务器尚未安装 XLSX 解析组件") from exc
    try:
        workbook = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    except Exception as exc:
        raise ImportFileError("XLSX 文件损坏或无法读取") from exc
    chapters = []
    try:
        for sheet in workbook.worksheets:
            lines = []
            for row in sheet.iter_rows(values_only=True):
                cells = [str(value).strip() for value in row if value not in (None, "")]
                if cells:
                    value = " ".join(cells)
                    total_chars += len(value) + 1
                    _check_budget(total_chars, max_chars)
                    lines.append(value)
            if lines:
                chapters.append({"title": sheet.title[:300], "text": "\n".join(lines)})
    finally:
        workbook.close()
    text = "\n\n".join(chapter["text"] for chapter in chapters)
    return text, chapters


def _extract_xls(data: bytes, max_chars: int) -> tuple[str, list[dict]]:
    total_chars = 0
    try:
        import xlrd
    except ImportError as exc:
        raise ImportFileError("服务器尚未安装 XLS 解析组件") from exc
    try:
        workbook = xlrd.open_workbook(file_contents=data, on_demand=True)
    except Exception as exc:
        raise ImportFileError("XLS 文件损坏或无法读取") from exc
    chapters = []
    try:
        for sheet in workbook.sheets():
            lines = []
            for row_index in range(sheet.nrows):
                cells = [
                    str(sheet.cell_value(row_index, column)).strip()
                    for column in range(sheet.ncols)
                    if str(sheet.cell_value(row_index, column)).strip()
                ]
                if cells:
                    value = " ".join(cells)
                    total_chars += len(value) + 1
                    _check_budget(total_chars, max_chars)
                    lines.append(value)
            if lines:
                chapters.append({"title": sheet.name[:300], "text": "\n".join(lines)})
    finally:
        workbook.release_resources()
    text = "\n\n".join(chapter["text"] for chapter in chapters)
    return text, chapters


def _extract_sqlite(data: bytes, max_chars: int) -> tuple[str, list[dict]]:
    if not data.startswith(b"SQLite format 3\x00"):
        raise ImportFileError("SQLite 文件头无效")
    path = ""
    try:
        with tempfile.NamedTemporaryFile(suffix=".sqlite", delete=False) as handle:
            handle.write(data)
            path = handle.name
        connection = sqlite3.connect(f"file:{path}?mode=ro", uri=True)
        try:
            connection.execute("PRAGMA query_only=ON")
            connection.execute("PRAGMA trusted_schema=OFF")
            table = connection.execute(
                "SELECT name FROM sqlite_master WHERE type='table' AND lower(name)='words'"
            ).fetchone()
            if not table:
                raise ImportFileError("SQLite 中没有 WORDS 单词表")
            columns = {
                str(row[1]).lower()
                for row in connection.execute('PRAGMA table_info("WORDS")').fetchall()
            }
            column = "stem" if "stem" in columns else "word" if "word" in columns else ""
            if not column:
                raise ImportFileError("WORDS 表中没有 stem 或 word 字段")
            if not re.fullmatch(r"[A-Za-z_][A-Za-z0-9_]*", column):
                raise ImportFileError("WORDS 表列名无效")
            cursor = connection.execute(
                f'SELECT "{column}" FROM "WORDS" WHERE "{column}" IS NOT NULL'
            )
            values = []
            total_chars = 0
            for row in cursor:
                value = str(row[0]).strip()
                if value:
                    total_chars += len(value) + (1 if values else 0)
                    _check_budget(total_chars, max_chars)
                    values.append(value)
        finally:
            connection.close()
    except sqlite3.DatabaseError as exc:
        raise ImportFileError("SQLite 数据库损坏或无法读取") from exc
    finally:
        if path:
            try:
                os.unlink(path)
            except OSError:
                pass
    text = "\n".join(values)
    return text, [{"title": "WORDS", "text": text}]


def extract_file_content(
    filename: str, data: bytes, max_chars: int
) -> tuple[str, str, list[dict]]:
    suffix = PurePosixPath(filename).suffix.lower()
    if suffix in {".txt", ".md"}:
        if len(data) > max_chars * 4 + 4 * 1024 * 1024:
            raise ImportFileError(f"正文过长，最多支持 {max_chars} 个字符")
        text = _decode_text(data)
        source_type = "file"
        chapters = split_text_chapters(text)
    elif suffix == ".epub":
        chapters = extract_epub_chapters(
            data,
            max_chapter_bytes=min(_ARCHIVE_UNCOMPRESSED_LIMIT, max_chars * 4 + 1024 * 1024),
            max_chars=max_chars,
        )
        text = "\n\n".join(chapter["text"] for chapter in chapters)
        source_type = "epub"
    elif suffix == ".pdf":
        text, chapters = _extract_pdf(data, max_chars)
        source_type = "pdf"
    elif suffix == ".docx":
        text, chapters = _extract_docx(data, max_chars)
        source_type = "docx"
    elif suffix == ".csv":
        text, chapters = _extract_csv(data, max_chars)
        source_type = "csv"
    elif suffix == ".xlsx":
        text, chapters = _extract_xlsx(data, max_chars)
        source_type = "xlsx"
    elif suffix == ".xls":
        text, chapters = _extract_xls(data, max_chars)
        source_type = "xls"
    elif suffix in {".db", ".sqlite"}:
        text, chapters = _extract_sqlite(data, max_chars)
        source_type = "sqlite"
    else:
        raise ImportFileError("支持 TXT、Markdown、PDF、DOCX、EPUB、CSV、XLSX、XLS、DB 和 SQLite")
    for chapter in chapters or []:
        chapter.setdefault(
            "paragraphs", _paragraphs_from_text(str(chapter.get("text") or ""))
        )
    text = text.strip()
    if not text:
        raise ImportFileError("文件内容为空")
    if len(text) > max_chars:
        raise ImportFileError(f"正文过长，最多支持 {max_chars} 个字符")
    return text, source_type, chapters


def _isolated_pdf_worker(connection, filename: str, data: bytes, max_chars: int) -> None:
    """在受限子进程中解析不可信 PDF，避免拖垮 Web 进程。"""
    try:
        try:
            import resource
            import sys

            resource.setrlimit(resource.RLIMIT_CPU, (20, 20))
            if sys.platform.startswith("linux"):
                memory_limit = 512 * 1024 * 1024
                resource.setrlimit(resource.RLIMIT_AS, (memory_limit, memory_limit))
        except (ImportError, OSError, ValueError):
            pass
        connection.send(("ok", extract_file_content(filename, data, max_chars)))
    except ImportFileError as exc:
        connection.send(("error", str(exc)))
    except Exception:
        connection.send(("error", "PDF 解析失败或超出资源限制"))
    finally:
        connection.close()


def extract_pdf_content_isolated(
    filename: str, data: bytes, max_chars: int, timeout_seconds: float = 30.0
) -> tuple[str, str, list[dict]]:
    """在独立进程解析 PDF，并施加 CPU、内存与墙钟上限。"""
    context = multiprocessing.get_context("spawn")
    parent, child = context.Pipe(duplex=False)
    process = context.Process(
        target=_isolated_pdf_worker,
        args=(child, filename, data, max_chars),
        daemon=True,
    )
    process.start()
    child.close()
    deadline = time.monotonic() + max(1.0, timeout_seconds)
    result = None
    while time.monotonic() < deadline:
        if parent.poll(0.1):
            result = parent.recv()
            break
        if not process.is_alive():
            break
    if process.is_alive():
        process.terminate()
    process.join(timeout=2)
    if process.is_alive():
        process.kill()
        process.join(timeout=1)
    parent.close()
    if not result:
        raise ImportFileError("PDF 解析超时或超出资源限制")
    status, payload = result
    if status != "ok":
        raise ImportFileError(str(payload))
    return payload
