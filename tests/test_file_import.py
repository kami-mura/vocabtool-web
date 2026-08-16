import gzip
import io
import sqlite3
import struct
import zipfile

import pytest
from fastapi import HTTPException

from app.api_support import _decode_upload_body
from app.file_import import (
    ImportFileError,
    _reject_xml_entity_declarations,
    extract_file_content,
)


def sample_epub(
    chapter_one: bytes = b"<html><body><h1>Chapter One</h1><p>Hello world.</p></body></html>",
) -> bytes:
    output = io.BytesIO()
    with zipfile.ZipFile(output, "w") as archive:
        archive.writestr("mimetype", "application/epub+zip")
        archive.writestr(
            "META-INF/container.xml",
            """<?xml version="1.0"?>
            <container xmlns="urn:oasis:names:tc:opendocument:xmlns:container">
              <rootfiles><rootfile full-path="OEBPS/content.opf"/></rootfiles>
            </container>""",
        )
        archive.writestr(
            "OEBPS/content.opf",
            """<?xml version="1.0"?>
            <package xmlns="http://www.idpf.org/2007/opf">
              <manifest>
                <item id="one" href="one.xhtml" media-type="application/xhtml+xml"/>
                <item id="two" href="two.xhtml" media-type="application/xhtml+xml"/>
                <item id="nav" href="nav.xhtml" media-type="application/xhtml+xml" properties="nav"/>
              </manifest>
              <spine><itemref idref="one"/><itemref idref="two"/></spine>
            </package>""",
        )
        archive.writestr("OEBPS/one.xhtml", chapter_one)
        archive.writestr("OEBPS/two.xhtml", "<html><body><h1>Chapter Two</h1><p>Keep reading.</p></body></html>")
        archive.writestr(
            "OEBPS/nav.xhtml",
            "<html><body><nav><a href='one.xhtml'>A New Beginning</a>"
            "<a href='two.xhtml'>The Second Road</a></nav></body></html>",
        )
    return output.getvalue()


def _lie_about_uncompressed_size(raw: bytes, target_name: str) -> bytes:
    """把指定成员在本地头与中央目录中声明的未压缩大小改为 1 字节。

    模拟攻击者伪造 zip 元数据：声明大小极小，实际解压内容巨大。
    """
    patched = bytearray(raw)
    with zipfile.ZipFile(io.BytesIO(raw)) as archive:
        info = archive.getinfo(target_name)
    struct.pack_into("<I", patched, info.header_offset + 22, 1)
    position = 0
    while True:
        position = patched.find(b"PK\x01\x02", position)
        if position < 0:
            break
        name_len = struct.unpack_from("<H", patched, position + 28)[0]
        extra_len = struct.unpack_from("<H", patched, position + 30)[0]
        comment_len = struct.unpack_from("<H", patched, position + 32)[0]
        name = bytes(patched[position + 46 : position + 46 + name_len]).decode()
        if name == target_name:
            struct.pack_into("<I", patched, position + 24, 1)
            break
        position += 46 + name_len + extra_len + comment_len
    return bytes(patched)


def test_epub_rejects_member_with_forged_uncompressed_size():
    """伪造声明大小的超大成员必须变成干净的导入错误，而不是 500 或读入内存。"""
    data = _lie_about_uncompressed_size(
        sample_epub(chapter_one=b"a" * 2_000_000), "OEBPS/one.xhtml"
    )
    with pytest.raises(ImportFileError):
        extract_file_content("book.epub", data, 10_000)


def test_archive_reader_counts_actual_uncompressed_bytes():
    from app.file_import import _ArchiveReader

    output = io.BytesIO()
    with zipfile.ZipFile(output, "w") as archive:
        archive.writestr("big.xml", b"a" * 2_000_000)
    data = output.getvalue()

    with zipfile.ZipFile(io.BytesIO(data)) as archive:
        reader = _ArchiveReader(archive, limit=1_000_000)
        with pytest.raises(ImportFileError, match="解压后过大"):
            reader.read(archive.infolist()[0])

    with zipfile.ZipFile(io.BytesIO(data)) as archive:
        reader = _ArchiveReader(archive, limit=10_000_000)
        content = reader.read(archive.infolist()[0])
        assert len(content) == 2_000_000
        assert reader.total == 2_000_000


def test_extract_epub_in_spine_order():
    text, source_type, _chapters = extract_file_content(
        "book.epub", sample_epub(), 10_000
    )
    assert source_type == "epub"
    assert "Chapter One" in text
    assert text.index("Chapter One") < text.index("Chapter Two")


def _epub_without_nav(chapter_raw: bytes) -> bytes:
    output = io.BytesIO()
    with zipfile.ZipFile(output, "w") as archive:
        archive.writestr("mimetype", "application/epub+zip")
        archive.writestr(
            "META-INF/container.xml",
            """<?xml version="1.0"?>
            <container xmlns="urn:oasis:names:tc:opendocument:xmlns:container">
              <rootfiles><rootfile full-path="OEBPS/content.opf"/></rootfiles>
            </container>""",
        )
        archive.writestr(
            "OEBPS/content.opf",
            """<?xml version="1.0"?>
            <package xmlns="http://www.idpf.org/2007/opf">
              <manifest>
                <item id="one" href="one.xhtml" media-type="application/xhtml+xml"/>
              </manifest>
              <spine><itemref idref="one"/></spine>
            </package>""",
        )
        archive.writestr("OEBPS/one.xhtml", chapter_raw)
    return output.getvalue()


@pytest.mark.parametrize("heading", [b"<h1></h1>", b"<h1>   </h1>"])
def test_epub_chapter_with_blank_heading_falls_back_to_filename_title(heading):
    data = _epub_without_nav(
        b"<html><body>" + heading + b"<p>Hello world.</p></body></html>"
    )
    text, source_type, chapters = extract_file_content("book.epub", data, 10_000)
    assert source_type == "epub"
    assert "Hello world." in text
    assert len(chapters) == 1
    assert chapters[0]["title"] == "one"
    assert "Hello world." in chapters[0]["text"]


def test_extract_epub_and_text_as_ordered_chapters():
    text, source_type, chapters = extract_file_content(
        "book.epub", sample_epub(), 10_000
    )
    assert source_type == "epub"
    assert [chapter["title"] for chapter in chapters] == [
        "A New Beginning",
        "The Second Road",
    ]
    assert "Keep reading" in chapters[1]["text"]
    assert "Chapter Two" in text


def test_decode_upload_body_honors_decompress_limit():
    raw = gzip.compress(b"x" * 1000)
    try:
        _decode_upload_body(
            None,
            raw,
            "book.json",
            upload_encoding="gzip",
            decompress_limit=100,
        )
    except HTTPException as exc:
        assert exc.status_code == 413
    else:
        raise AssertionError("expected 413 for over-limit decompression")

    data, _filename = _decode_upload_body(
        None,
        raw,
        "book.json",
        upload_encoding="gzip",
        decompress_limit=2000,
    )
    assert len(data) == 1000

    text, source_type, chapters = extract_file_content(
        "book.md",
        b"# Chapter One\nHello.\n\n# Chapter Two\nWorld.",
        10_000,
    )
    assert source_type == "file"
    assert [chapter["title"] for chapter in chapters] == [
        "Chapter One",
        "Chapter Two",
    ]


def test_rejects_xml_entity_declarations_before_parsing():
    with pytest.raises(ImportFileError, match="XML 实体声明"):
        _reject_xml_entity_declarations(
            b'<?xml version="1.0"?><!DOCTYPE package [<!ENTITY x "boom">]><package/>',
            "测试 XML",
        )
    _reject_xml_entity_declarations(b"<package/>", "测试 XML")


def test_docx_rejects_doctype_xml_before_library_parse():
    output = io.BytesIO()
    with zipfile.ZipFile(output, "w") as archive:
        archive.writestr(
            "[Content_Types].xml",
            '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types"/>',
        )
        archive.writestr(
            "word/document.xml",
            '<!DOCTYPE document [<!ENTITY x "boom">]><w:document/>',
        )
    with pytest.raises(ImportFileError, match="XML 实体声明"):
        extract_file_content("bad.docx", output.getvalue(), 10_000)


def test_epub_stops_when_extracted_text_exceeds_budget():
    with pytest.raises(ImportFileError, match="正文过长"):
        extract_file_content("book.epub", sample_epub(), 10)


def test_extract_csv_and_sqlite_word_sources():
    text, source_type, chapters = extract_file_content(
        "words.csv",
        "word,meaning\nrun,跑步\nadaptive,适应性强的".encode(),
        10_000,
    )
    assert source_type == "csv"
    assert "run 跑步" in text
    assert chapters[0]["title"] == "CSV 数据"

    connection = sqlite3.connect(":memory:")
    try:
        connection.execute("CREATE TABLE WORDS (stem TEXT)")
        connection.executemany(
            "INSERT INTO WORDS(stem) VALUES (?)", [("run",), ("adaptive",)]
        )
        data = connection.serialize()
    finally:
        connection.close()
    text, source_type, chapters = extract_file_content(
        "words.sqlite", data, 10_000
    )
    assert source_type == "sqlite"
    assert text.splitlines() == ["run", "adaptive"]
    assert chapters == [
        {
            "title": "WORDS",
            "text": "run\nadaptive",
            "paragraphs": [
                {"kind": "p", "html": "run<br>adaptive", "text": "run\nadaptive"}
            ],
        }
    ]


def test_extract_docx_and_xlsx_with_named_sections():
    from docx import Document
    from openpyxl import Workbook

    document = Document()
    document.add_heading("First Chapter", level=1)
    document.add_paragraph("People run every morning.")
    document.add_heading("Second Chapter", level=1)
    document.add_paragraph("Adaptive systems improve.")
    docx = io.BytesIO()
    document.save(docx)
    text, source_type, chapters = extract_file_content(
        "book.docx", docx.getvalue(), 10_000
    )
    assert source_type == "docx"
    assert "People run every morning" in text
    assert [chapter["title"] for chapter in chapters] == [
        "First Chapter",
        "Second Chapter",
    ]

    workbook = Workbook()
    first = workbook.active
    first.title = "Vocabulary"
    first.append(["run", "跑步"])
    second = workbook.create_sheet("Sentences")
    second.append(["People run every morning."])
    xlsx = io.BytesIO()
    workbook.save(xlsx)
    workbook.close()
    text, source_type, chapters = extract_file_content(
        "book.xlsx", xlsx.getvalue(), 10_000
    )
    assert source_type == "xlsx"
    assert "People run every morning" in text
    assert [chapter["title"] for chapter in chapters] == [
        "Vocabulary",
        "Sentences",
    ]


def test_epub_rejects_oversized_single_chapter_before_parsing():
    """单章超过 4MB 绝对上限必须在 HTML 解析之前被拒：
    之前单章可以吃满 max_chapter_bytes 并先完成整章解析才查字符预算，
    一个压缩包就能在 Web 进程内放大出数百 MB 内存峰值。"""
    big = b"a" * (4 * 1024 * 1024 + 1)
    data = sample_epub(chapter_one=big)
    # 总字符预算足够大（1000 万），拒绝只能来自单章绝对上限。
    with pytest.raises(ImportFileError, match="解压后过大"):
        extract_file_content("book.epub", data, 10_000_000)


def test_epub_normal_chapters_still_import_within_budget():
    """修复不应影响正常大小的书：预算内多章照常解析。"""
    text, source_type, _chapters = extract_file_content(
        "book.epub", sample_epub(), 10_000
    )
    assert source_type == "epub"
    assert "Chapter One" in text and "Chapter Two" in text


def test_xml_entity_rejection_catches_utf16_documents():
    """UTF-16 编码的 DOCTYPE/ENTITY 不能绕过字节子串检测：
    转码后必须同样被拒（见 docs/审查整改清单.md P2-13）。"""
    malicious = (
        '<?xml version="1.0" encoding="UTF-16"?>\n'
        '<!DOCTYPE root [<!ENTITY xxe SYSTEM "file:///etc/passwd">]>\n'
        "<root>&xxe;</root>"
    )
    for encoding in ("utf-16", "utf-16-le", "utf-16-be"):
        encoded = malicious.encode(encoding)
        assert b"<!DOCTYPE" not in encoded  # 前置条件：ASCII 子串确实搜不到
        with pytest.raises(ImportFileError, match="不安全的 XML 实体声明"):
            _reject_xml_entity_declarations(encoded, "DOCX")

    # 无 BOM 但声明了 UTF-16 的变体同样要拦下。
    import codecs as _codecs

    no_bom = malicious.encode("utf-16-le")  # utf-16-le 编码本身不产生 BOM
    assert not no_bom.startswith(_codecs.BOM_UTF16_LE)
    with pytest.raises(ImportFileError, match="不安全的 XML 实体声明"):
        _reject_xml_entity_declarations(no_bom, "DOCX")

    # 正常的 UTF-16 文档（无实体声明）不受影响。
    clean = (
        '<?xml version="1.0" encoding="UTF-16"?>\n<root><child>ok</child></root>'
    ).encode("utf-16")
    _reject_xml_entity_declarations(clean, "DOCX")  # 不应抛错
